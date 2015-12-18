####################################################
#                  Revision: 1.0                   #
#              Updated on: 11/06/2015              #
####################################################

####################################################
#                                                  #
#   This File appends user specified number of     #
#   extracted/modified files to the Modified       #
#   (i.e. Keywords replaced) Word Template.        #
#                                                  #
#   Author: Zankar Sanghavi                        #
#                                                  #
#   © Dot Hill Systems Corporation                 #
#                                                  #
####################################################


###################################################
#
#   Importing required packages
# 
###################################################
import pandas
import csv
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm
import re
import zipfile
import time
import os
from openpyxl import load_workbook

import fixed_data
import report_functions
import extract_lists
import user_inputs
import modify_word_docx

###################################################
#
#   This Function will open a Modified Word Template
#   (i.e. Word Template with Keywords replaced) and 
#   add user required files at the bottom of the 
#   Template.
# 
###################################################
def Append_Files(file_names, num, fixed_dir, part_no,
                   model_no_names, cap_names, chassis_names
                   , cntrllr_names, fw_type_temp, fw_no_names
                   , test_name):
   
    rf = report_functions.Report_Functions
    
    document = Document(r''+str(fixed_dir)+'\\'+str(part_no)
                +str(test_name)+'.docx')
    document.add_page_break()


    ###################################################
    #
    #   A loop to append/add multiple files to Modified
    #   Word template 
    # 
    ###################################################
    for f in range(int(num)):
    
        progress=(round((float(100/int(num))*f),2)) 
        # to show Report Progress
        print('\nReport Progess: ',progress,'%\n')
        
        m_data = pandas.read_csv(open(r''+file_names[f]
                +'_Modified.csv'),header=None)
        md=np.array(m_data)
        
        
        ###################################################
        #
        #  Converting numbers to 1 decimal place
        # 
        ###################################################
        for j in range(4,8):
            for i in range(1,len(md[:,1])):
            
                temp=md[i,j]
                precise_d=float(temp)
                precise_d=round(precise_d,1)
                md[i,j]=str(precise_d)
                
                
        ###################################################
        #
        #  Appending files from the section of Modified 
        #  word template with all Formatting. 
        # 
        ###################################################
        document.add_heading('Performance Results - chassis ' +str(f+1)
                + '\nModel: '+model_no_names[f]+ ', ' +cap_names[f]  
                + ' in ' +chassis_names[f]+ '/' +cntrllr_names[f]
                + ' chassis, HDD '+str(fw_type_temp)+' FW: '
                +fw_no_names[f],level=3)


        section = document.sections[-1] 
        # last section in document
        section.start_type = WD_SECTION.NEW_PAGE
        
        section.top_margin= Inches(0.3)
        section.bottom_margin=Inches(0)

        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        r=len(md[:,0])
        c=len(md[0,:])

        table_style = document.styles["Normal"]
        table_font=table_style.font
        table = document.add_table(rows=r,cols=c)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        table.autofit=False
        rf.set_column_width(table.columns[0], Cm(1.3)) #alignment
        rf.set_column_width(table.columns[1], Cm(2.3)) #drive
        rf.set_column_width(table.columns[2], Cm(3.8)) #target name
        rf.set_column_width(table.columns[3], Cm(2.8)) #acess spec
        rf.set_column_width(table.columns[4], Cm(1.5)) #IOps
        rf.set_column_width(table.columns[5], Cm(1.4)) #MBps
        rf.set_column_width(table.columns[6], Cm(1.6)) #Avg. Latency
        rf.set_column_width(table.columns[7], Cm(1.6)) #Max. Latency
        rf.set_column_width(table.columns[8], Cm(1.1)) #QD
        rf.set_column_width(table.columns[9], Cm(1.3)) #Read Errors
        rf.set_column_width(table.columns[10], Cm(1.3)) #Write Errors

        table.style = 'Table Grid'

        for i in range(r):
            for j in range(c):
                table_font.size=Pt(9)

                hdr_cells=table.rows[i].cells
                if i == 0:
                    hdr_cells[j].text = md[i,j]

                else:

                    hdr_cells[j].text = md[i,j]  
                

        document.add_picture(r''+file_names[f]
            +'_Modified.csv_Plot_1.png',width=Inches(8)
            ,height=Inches(9))
            
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_picture(r''+file_names[f]
            +'_Modified.csv_Plot_2.png'
            ,width=Inches(7.5),height=Inches(9.3))
            
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.save(r''+str(fixed_dir)+'\\'+str(part_no)
            +str(test_name)+'.docx')
            
    # remove temporary generated document    
    os.remove(r''+str(fixed_dir)+'\\temp_doc.docx')   
    os.remove(r''+file_names[f]+'_Modified.csv')
    os.remove(r''+file_names[f]+'_Modified.csv_Plot_1.png')
    os.remove(r''+file_names[f]+'_Modified.csv_Plot_2.png')
    
    print('\nYour Report is ready!\n')
    
    
#####################################
#              END                  #
#####################################
